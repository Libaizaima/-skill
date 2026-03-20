# -*- coding: utf-8 -*-
"""DOCX 报告生成模块"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Dict, Any, List
import os


def generate(results: Dict[str, Any], output_path: str):
    """
    根据分析结果生成 DOCX 报告。

    Args:
        results: analyzer.analyze_all() 返回的分析结果字典
        output_path: 输出文件路径
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    company_name = results.get('company_name', '未知公司')
    overall = results.get('overall', {})

    # ==================== 报告标题 ====================
    _add_centered_title(doc, '客户分析报告', size=18, bold=True)
    _add_centered_title(doc, company_name, size=14, bold=True)
    doc.add_paragraph('')

    # ==================== 一、个人情况 ====================
    _add_heading(doc, '一、个人情况')

    _add_paragraph(doc, '①个人以及家庭信息（籍贯、年龄、居住情况）')
    _add_paragraph(doc, '（待补充）', italic=True)

    _add_paragraph(doc, '②个人征信情况（负债、逾期、贷款到期预警以及查询情况）：')
    credit_data = results.get('credit_data', {})
    personal_credit = credit_data.get('personal', None) if credit_data else None
    company_credit = credit_data.get('company', None) if credit_data else None

    _add_personal_credit_table(doc, personal_credit)
    ai = results.get('ai_analysis', {})
    if ai.get('credit_personal'):
        _add_paragraph(doc, ai['credit_personal'])
    elif personal_credit:
        _add_personal_credit_analysis(doc, personal_credit)
    else:
        _add_paragraph(doc, '（征信分析待补充）', italic=True)

    _add_paragraph(doc, '企业征信：')
    _add_company_credit_table(doc, company_credit)
    if ai.get('credit_company'):
        _add_paragraph(doc, ai['credit_company'])
    elif company_credit:
        _add_company_credit_analysis(doc, company_credit)
    else:
        _add_paragraph(doc, '（企业征信分析待补充）', italic=True)

    doc.add_paragraph('')

    # ==================== ②在职员工/发薪 ====================
    _add_heading(doc, '二、在职员工（发薪情况）')

    salary_df = results.get('salary', pd.DataFrame())
    if not salary_df.empty:
        _add_sub_heading(doc, '发薪情况（月度统计）')
        _add_table(doc, salary_df, ['月份', '发薪笔数', '发薪金额（元）'])

        total_salary = salary_df['发薪金额'].sum()
        avg_salary = total_salary / len(salary_df) if len(salary_df) > 0 else 0
        months_text = f"{salary_df.iloc[0]['月份']}至{salary_df.iloc[-1]['月份']}" if len(salary_df) > 0 else ''
        _add_paragraph(doc,
            f"发薪分析：{months_text}流水中识别到工资类支出合计约{total_salary:,.2f}元，"
            f"月均约{avg_salary:,.2f}元，显示企业存在持续发薪行为。")
    else:
        _add_paragraph(doc, '未在流水中识别到工资类支出记录。', italic=True)

    doc.add_paragraph('')

    # ==================== ③经营场地/交租 ====================
    _add_heading(doc, '三、经营场地（交租情况）')

    # 场地信息空白模板
    _add_table(doc, pd.DataFrame(), ['合同签订周期', '租期\n(年）', '场地面积\n（平方米）', '租金', '出租方', '交租情况'],
               empty_rows=1)

    rent_df = results.get('rent', pd.DataFrame())
    if not rent_df.empty:
        _add_sub_heading(doc, '交租情况（月度统计）')
        _add_table(doc, rent_df, ['月份', '交租笔数', '交租金额（元）'])

        total_rent = rent_df['交租金额'].sum()
        avg_rent = total_rent / len(rent_df) if len(rent_df) > 0 else 0
        _add_paragraph(doc,
            f"交租分析：流水中识别到租金类支出合计约{total_rent:,.2f}元，"
            f"月均约{avg_rent:,.2f}元。")
    else:
        _add_paragraph(doc, '未在流水中识别到明确的租金支出记录。', italic=True)

    doc.add_paragraph('')

    # ==================== ④对公流水分析 ====================
    _add_heading(doc, '四、对公流水分析')

    if overall:
        _add_sub_heading(doc, '总体结论')
        _add_paragraph(doc,
            f"1. 全年总收入 {overall['total_income']:,.2f} 元，"
            f"总支出 {overall['total_expense']:,.2f} 元，"
            f"净额 {overall['net']:,.2f} 元。")
        _add_paragraph(doc,
            f"2. 其中经营性有效流水：收入以下游客户回款为主，支出以上游采购货款为主。")
        _add_paragraph(doc,
            f"3. 全年交易笔数 {overall['total_count']} 笔；"
            f"经营性有效流水 {overall['biz_count']} 笔；"
            f"非经营性/特殊流水 {overall['other_count']} 笔。")
        _add_paragraph(doc,
            f"4. 账户余额最低 {overall['balance_min']:,.2f} 元，"
            f"最高 {overall['balance_max']:,.2f} 元，"
            f"末笔余额 {overall['balance_last']:,.2f} 元。")

        # AI 流水分析
        if ai.get('flow_analysis'):
            _add_sub_heading(doc, 'AI 分析')
            _add_paragraph(doc, ai['flow_analysis'])

    # 结息统计
    interest_df = results.get('interest', pd.DataFrame())
    if not interest_df.empty:
        _add_sub_heading(doc, '结息统计表')
        _add_table(doc, interest_df, ['交易时间', '对方户名', '收入金额', '支出金额', '摘要'])

    # 流水分类汇总
    category_df = results.get('category_summary', pd.DataFrame())
    if not category_df.empty:
        _add_sub_heading(doc, '各类型流水汇总表')
        # 格式化金额
        display_df = category_df.copy()
        display_df['收入'] = display_df['收入'].apply(lambda x: f'{x:,.2f}')
        display_df['支出'] = display_df['支出'].apply(lambda x: f'{x:,.2f}')
        display_df['净额'] = display_df['净额'].apply(lambda x: f'{x:,.2f}')
        _add_table(doc, display_df, ['分类', '收入', '支出', '净额', '笔数'])

    # 月度汇总表
    monthly_df = results.get('monthly', pd.DataFrame())
    if not monthly_df.empty:
        _add_sub_heading(doc, '月度收支汇总表')
        _add_table(doc, monthly_df, ['月份', '收入（万元）', '支出（万元）', '净流入（万元）'])

    # 前10收入对手方
    top_income = results.get('top_income', pd.DataFrame())
    if not top_income.empty:
        _add_sub_heading(doc, '经营性收入前十对象')
        display_income = top_income.copy()
        display_income['经营收入合计'] = display_income['经营收入合计'].apply(lambda x: f'{x:,.2f}')
        _add_table(doc, display_income, ['对方户名', '经营收入合计'])

    # 前10支出对手方
    top_expense = results.get('top_expense', pd.DataFrame())
    if not top_expense.empty:
        _add_sub_heading(doc, '经营性支出前十对象')
        display_expense = top_expense.copy()
        display_expense['经营支出合计'] = display_expense['经营支出合计'].apply(lambda x: f'{x:,.2f}')
        _add_table(doc, display_expense, ['对方户名', '经营支出合计'])

    doc.add_paragraph('')

    # ==================== 六、应收分析 ====================
    _add_heading(doc, '五、应收分析')

    recv_total = results.get('receivable_total', {})
    total_wan = recv_total.get('总应收(万元)', 0)
    currency = recv_total.get('币种', 'CNY')
    rate = recv_total.get('汇率', 1)
    total_foreign = recv_total.get('总应收(外币)', 0)

    if total_wan > 0:
        text = f"应收明细账显示，应收合计约{total_wan:,.2f}万元"
        if currency != 'CNY' and total_foreign > 0:
            text += f"（外币{total_foreign:,.2f} {currency}，预算汇率{rate}）"
        text += "。"
        _add_paragraph(doc, text)

    recv_top = results.get('receivable_top', pd.DataFrame())
    if not recv_top.empty:
        # 动态确定要显示的列
        display_cols = ['客户名称']
        if '币种' in recv_top.columns:
            display_cols.append('币种')
        if '预算汇率' in recv_top.columns:
            display_cols.append('预算汇率')
        if '合计(外币)' in recv_top.columns:
            display_cols.append('合计(外币)')
        if '合计(万元)' in recv_top.columns:
            display_cols.append('合计(万元)')
        # 添加月度列
        for col in recv_top.columns:
            if col not in display_cols and col not in ('序号',):
                display_cols.append(col)
        _add_table(doc, recv_top, display_cols)

    doc.add_paragraph('')

    # ==================== 七、应付分析 ====================
    _add_heading(doc, '六、应付分析')

    pay_total = results.get('payable_total', {})
    total_pay_wan = pay_total.get('总应付(万元)', 0)
    if total_pay_wan > 0:
        _add_paragraph(doc,
            f"应付明细表显示，供应商应付余额合计约{total_pay_wan:,.2f}万元。")

    pay_top = results.get('payable_top', pd.DataFrame())
    if not pay_top.empty:
        # 动态确定要显示的列
        display_cols = ['供应商']
        if '合计(万元)' in pay_top.columns:
            display_cols.append('合计(万元)')
        # 添加月度列
        for col in pay_top.columns:
            if col not in display_cols:
                display_cols.append(col)
        _add_table(doc, pay_top, display_cols)

    doc.add_paragraph('')

    # ==================== ⑨开票情况统计分析 ====================
    _add_heading(doc, '七、开票情况统计分析')

    # 进项票年度
    inv_in_yearly = results.get('invoice_in_yearly', pd.DataFrame())
    if not inv_in_yearly.empty:
        _add_sub_heading(doc, '进项票年度统计表')
        _add_table(doc, inv_in_yearly, ['年份', '开票张数', '开票金额（万元）'])

    # 销项票年度
    inv_out_yearly = results.get('invoice_out_yearly', pd.DataFrame())
    if not inv_out_yearly.empty:
        _add_sub_heading(doc, '销项票年度统计表')
        _add_table(doc, inv_out_yearly, ['年份', '开票张数', '开票金额（万元）'])

    # 进销项对比
    inv_comp = results.get('invoice_comparison', pd.DataFrame())
    if not inv_comp.empty:
        _add_sub_heading(doc, '进销项年度对比表')
        _add_table(doc, inv_comp, ['年份', '进项金额（万元）', '销项金额（万元）', '销-进差额（万元）'])

    # 进项月度
    inv_in_monthly = results.get('invoice_in_monthly', pd.DataFrame())
    if not inv_in_monthly.empty:
        _add_sub_heading(doc, '进项票月度统计表')
        _add_table(doc, inv_in_monthly, ['月份', '开票张数', '开票金额（万元）'])

    # 销项月度
    inv_out_monthly = results.get('invoice_out_monthly', pd.DataFrame())
    if not inv_out_monthly.empty:
        _add_sub_heading(doc, '销项票月度统计表')
        _add_table(doc, inv_out_monthly, ['月份', '开票张数', '开票金额（万元）'])

    # 出口发票统计
    export_stats = results.get('export_stats', {})
    if export_stats.get('has_export'):
        _add_sub_heading(doc, '出口发票统计（按币种）')
        export_summary = export_stats.get('summary', [])
        if export_summary:
            export_df = pd.DataFrame(export_summary)
            # 格式化
            display_export = export_df.copy()
            display_export['外币金额合计'] = display_export['外币金额合计'].apply(lambda x: f'{x:,.2f}')
            display_export['RMB金额合计'] = display_export['RMB金额合计'].apply(lambda x: f'{x:,.2f}')
            display_export['RMB金额合计（万元）'] = export_df['RMB金额合计'].apply(lambda x: f'{x/10000:,.2f}')
            _add_table(doc, display_export, ['币种', '张数', '外币金额合计', 'RMB金额合计', 'RMB金额合计（万元）', '平均汇率'])

        ratio = export_stats.get('export_ratio', 0)
        total_rmb = export_stats.get('total_rmb', 0)
        total_all = export_stats.get('total_rmb_all', 0)
        _add_paragraph(doc,
            f"出口发票分析：出口发票RMB折算金额合计{total_rmb:,.2f}元"
            f"（{total_rmb/10000:,.2f}万元），"
            f"占全部销项票金额的{ratio}%。"
            f"全部销项票金额合计{total_all:,.2f}元（{total_all/10000:,.2f}万元）。")

    doc.add_paragraph('')

    # AI 发票分析
    if ai.get('invoice_analysis'):
        _add_sub_heading(doc, '发票分析')
        _add_paragraph(doc, ai['invoice_analysis'])

    doc.add_paragraph('')

    # ==================== 上下游客户分析 ====================
    _add_heading(doc, '八、上下游客户分析')
    _add_paragraph(doc, '①公司主要客户占比分析（包含合作时长、回款账期、应收应付）：')

    # 基于前5大客户和供应商生成框架
    if not top_income.empty:
        for idx, (_, row) in enumerate(top_income.head(3).iterrows(), 1):
            name = row['对方户名']
            amount = row['经营收入合计']
            role = '销货方' if idx <= 2 else '供货方'
            _add_paragraph(doc,
                f"{idx}、{name}（{role}）：经营性收入合计约{amount}元。"
                f"（合作时长、交货周期、核心竞争力、负面信息待补充）", italic=True)

    doc.add_paragraph('')

    # ==================== 下游风险 ====================
    _add_heading(doc, '九、前五大下游客户公开案件/执行风险汇总')
    if not top_income.empty:
        risk_data = []
        for _, row in top_income.head(5).iterrows():
            risk_data.append({
                '下游客户': row['对方户名'],
                '经营收入金额': row['经营收入合计'],
                '公开风险摘要': '（待补充）',
                '初步判断': '（待补充）',
            })
        risk_df = pd.DataFrame(risk_data)
        _add_table(doc, risk_df, ['下游客户', '经营收入金额', '公开风险摘要', '初步判断'])

    doc.add_paragraph('')

    # ==================== 风险评估 ====================
    if ai.get('risk_assessment'):
        _add_heading(doc, '十、综合风险评估')
        _add_paragraph(doc, ai['risk_assessment'])
        doc.add_paragraph('')

    # ==================== 总结 ====================
    _add_heading(doc, '十一、总结')
    if ai.get('summary'):
        _add_paragraph(doc, ai['summary'])
    else:
        _add_paragraph(doc, '优点：（待根据分析结果补充）', italic=True)
        _add_paragraph(doc, '劣势：（待根据分析结果补充）', italic=True)
        _add_paragraph(doc, '考察需要补充的资料：（待补充）', italic=True)
        _add_paragraph(doc, '投资建议：（待根据分析结果补充）', italic=True)

    # 保存
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
    doc.save(output_path)
    print(f"[INFO] 报告已生成: {output_path}")


# ==================== 辅助函数 ====================

def _add_centered_title(doc, text, size=18, bold=True):
    """添加居中标题"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def _add_heading(doc, text):
    """添加一级标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def _add_sub_heading(doc, text):
    """添加二级标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def _add_paragraph(doc, text, bold=False, italic=False):
    """添加正文段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True


def _add_table(doc, df: pd.DataFrame, columns: List[str], empty_rows: int = 0):
    """
    添加表格到文档。

    Args:
        doc: Document 对象
        df: 数据 DataFrame
        columns: 要显示的列名列表（也作为表头）
        empty_rows: 额外添加的空行数
    """
    num_rows = len(df) + 1 + empty_rows  # +1 for header
    num_cols = len(columns)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表头
    header_row = table.rows[0]
    for ci, col_name in enumerate(columns):
        cell = header_row.cells[ci]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(col_name)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 表头背景色 - 浅灰
        _set_cell_bg(cell, 'D9E2F3')

    # 填充数据
    if not df.empty:
        for ri, (_, row) in enumerate(df.iterrows()):
            data_row = table.rows[ri + 1]
            for ci, col_name in enumerate(columns):
                # 尝试匹配列名（可能列名不完全一致）
                val = _get_cell_value(row, col_name, columns, df.columns)
                cell = data_row.cells[ci]
                cell.text = ''
                para = cell.paragraphs[0]
                run = para.add_run(str(val))
                run.font.size = Pt(9)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                # 数值右对齐
                if _is_numeric_str(str(val)):
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return table


def _get_cell_value(row, display_col: str, display_columns: List[str], df_columns) -> str:
    """获取单元格的值，处理列名映射"""
    # 精确匹配
    if display_col in row.index:
        val = row[display_col]
        return _format_value(val, display_col)

    # 模糊匹配: 先用display_col中去掉括号内容的部分匹配
    base_name = display_col.split('（')[0].split('(')[0].strip()
    for col in row.index:
        if base_name in str(col):
            val = row[col]
            return _format_value(val, display_col)

    return ''


def _format_value(val, col_name: str) -> str:
    """格式化单元格值"""
    if pd.isna(val) or val is None:
        return ''
    if isinstance(val, float):
        if '万元' in col_name or '余额' in col_name:
            return f'{val:,.2f}'
        if '金额' in col_name:
            return f'{val:,.2f}'
        if val == int(val):
            return str(int(val))
        return f'{val:.2f}'
    return str(val)


def _is_numeric_str(s: str) -> bool:
    """判断字符串是否像数字"""
    s = s.replace(',', '').replace('，', '').strip()
    try:
        float(s)
        return True
    except ValueError:
        return False


def _set_cell_bg(cell, color_hex: str):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_personal_credit_table(doc, data):
    """添加个人征信表格（有数据则填充，无数据则空白）"""
    p = data or {}
    rows = [
        ['征信时间：', p.get('征信时间', ''), '姓名：', p.get('姓名', '')],
        ['近一个月查询次数', p.get('近1月查询次数', ''), '近三个月查询次数', p.get('近3月查询次数', '')],
        ['信用卡总额度：', p.get('信用卡总额度', ''), '已用额度：', p.get('信用卡已用额度', '')],
        ['贷款户数：', p.get('贷款账户数', ''), '月供金额：', p.get('月供金额', '')],
        ['贷款总额度：', p.get('贷款总额度', ''), '未还总本金', p.get('贷款未还总本金', '')],
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = 'Table Grid'
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri].cells[ci]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if ci % 2 == 0:
                run.font.bold = True


def _add_personal_credit_analysis(doc, p):
    """根据个人征信数据生成分析文字"""
    lines = []
    name = p.get('姓名', '')
    card_count = p.get('信用卡账户数', '0')
    card_open = p.get('信用卡未结清数', '0')
    card_overdue = p.get('信用卡逾期账户数', '--')
    loan_count = p.get('贷款账户数', '0')
    loan_open = p.get('贷款未结清数', '0')
    loan_overdue = p.get('贷款逾期账户数', '--')
    total_limit = p.get('贷款总额度', '')
    total_balance = p.get('贷款未还总本金', '')
    q1 = p.get('近1月查询次数', '0')
    q3 = p.get('近3月查询次数', '0')
    q6 = p.get('近6月查询次数', '0')

    lines.append(f"征信分析：{name}共有信用卡{card_count}张（未结清{card_open}张），"
                 f"贷款{loan_count}笔（未结清{loan_open}笔）。")

    if card_overdue == '--' and loan_overdue == '--':
        lines.append("无逾期记录，征信整体良好。")
    elif card_overdue != '--' or loan_overdue != '--':
        lines.append(f"存在逾期记录（信用卡逾期{card_overdue}户，贷款逾期{loan_overdue}户），需关注。")

    if total_limit and total_balance:
        lines.append(f"贷款总额度{total_limit}元，未还总本金{total_balance}元。")

    lines.append(f"近1月查询{q1}次，近3月{q3}次，近6月{q6}次。")

    # 贷款明细
    loan_details = p.get('贷款明细', [])
    if loan_details:
        lines.append('贷款明细：')
        for loan in loan_details:
            status_text = loan.get('状态', '')
            limit = loan.get('额度', '')
            balance = loan.get('余额', '')
            detail = f"  - {loan.get('机构', '')}" 
            if limit:
                detail += f"，额度{limit}元"
            if balance:
                detail += f"，余额{balance}元"
            if status_text:
                detail += f"，{status_text}"
            lines.append(detail)

    for line in lines:
        _add_paragraph(doc, line)


def _add_company_credit_table(doc, data):
    """添加企业征信表格"""
    c = data or {}
    company_name = c.get('企业名称', '')
    rows = [
        ['征信时间：', c.get('征信时间', ''), '公司：', company_name],
        ['借贷余额：', c.get('借贷余额', ''), '担保余额：', c.get('担保余额', '')],
        ['信贷机构数：', c.get('信贷机构数', ''), '未结清机构数：', c.get('未结清机构数', '')],
    ]

    # 如果有贷款明细，添加表头 + 明细行
    loan_details = c.get('贷款明细', [])
    has_loans = len(loan_details) > 0
    if has_loans:
        rows.append(['贷款机构', '批款额度（万）', '未还本金（万）', '状态'])
        for loan in loan_details:
            limit = loan.get('额度', '')
            balance = loan.get('余额', '')
            try:
                limit_wan = f"{float(limit) / 10000:.2f}" if limit else ''
            except:
                limit_wan = limit
            try:
                bal_wan = f"{float(balance) / 10000:.2f}" if balance else ''
            except:
                bal_wan = balance
            rows.append([loan.get('机构', ''), limit_wan, bal_wan, loan.get('状态', '')])
    else:
        # 空白行
        rows.append(['贷款机构', '批款额度（万）', '未还本金（万）', '使用周期'])
        rows.append(['', '', '', ''])

    table = doc.add_table(rows=len(rows), cols=4)
    table.style = 'Table Grid'
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri].cells[ci]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 表头行和标签列加粗
            header_row_idx = 3  # 贷款明细表头行
            if ri == header_row_idx or (ri < header_row_idx and ci % 2 == 0):
                run.font.bold = True
                if ri == header_row_idx:
                    _set_cell_bg(cell, 'D9E2F3')


def _add_company_credit_analysis(doc, c):
    """根据企业征信数据生成分析文字"""
    name = c.get('企业名称', '该企业')
    inst_count = c.get('信贷机构数', '0')
    open_count = c.get('未结清机构数', '0')
    borrow = c.get('借贷余额', '0')
    guarantee = c.get('担保余额', '0')

    if inst_count == '0' and borrow == '0':
        _add_paragraph(doc, f"企业征信分析：{name}当前无信贷交易记录，借贷余额为零，担保余额为零，企业征信干净。")
    else:
        _add_paragraph(doc, f"企业征信分析：{name}当前有{inst_count}家信贷机构交易记录，"
                       f"未结清{open_count}家。借贷余额{borrow}元，担保余额{guarantee}元。")

    query_count = len(c.get('查询记录', []))
    if query_count > 0:
        _add_paragraph(doc, f"企业征信查询记录{query_count}条。")

